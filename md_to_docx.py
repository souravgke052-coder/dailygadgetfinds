#!/usr/bin/env python3
"""
Convert DailyGadgetFinds_MOP.md to DailyGadgetFinds_MOP.docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = "/Users/souravlobra/Projects/Blogs-Websites/mobile_accessories-blog/DailyGadgetFinds_MOP.md"
OUTPUT_FILE = "/Users/souravlobra/Projects/Blogs-Websites/mobile_accessories-blog/DailyGadgetFinds_MOP.docx"


def set_cell_shading(cell, color):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_paragraph_with_style(doc, text, style_name=None, bold=False, italic=False,
                              font_name="Calibri", font_size=11, gray_bg=False):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style_name)

    if gray_bg:
        # Add shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
        pPr.append(shd)

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)

    return p


def parse_inline_formatting(text):
    """Parse inline bold (**text**) and inline code (`code`) from text."""
    parts = []
    # Combined pattern for bold and inline code
    pattern = r'\*\*(.+?)\*\*|`([^`]+)`'
    last_end = 0

    for match in re.finditer(pattern, text):
        # Add plain text before match
        if match.start() > last_end:
            parts.append(("plain", text[last_end:match.start()]))
        if match.group(1):  # Bold
            parts.append(("bold", match.group(1)))
        elif match.group(2):  # Inline code
            parts.append(("code", match.group(2)))
        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        parts.append(("plain", text[last_end:]))

    return parts


def add_formatted_paragraph(doc, text, style_name=None, font_name="Calibri", font_size=11, gray_bg=False):
    """Add a paragraph with inline formatting preserved."""
    p = doc.add_paragraph(style=style_name)

    if gray_bg:
        pPr = p._p.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
        pPr.append(shd)

    parts = parse_inline_formatting(text)

    for part_type, part_text in parts:
        run = p.add_run(part_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        if part_type == "bold":
            run.bold = True
        elif part_type == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            # Light gray background for inline code
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
            rPr.append(shd)

    return p


def parse_table_md(table_text):
    """Parse markdown table into list of lists."""
    lines = table_text.strip().split('\n')
    rows = []

    for line in lines:
        # Skip separator line (|---|)
        if re.match(r'^\|[\s\-:]+\|', line):
            continue
        # Extract cells (handle | at start and end)
        cells = re.findall(r'\|([^|]+)', line)
        cells = [cell.strip() for cell in cells]
        if cells:
            rows.append(cells)

    return rows


def add_table_to_doc(doc, rows):
    """Add a markdown table to the document with proper formatting."""
    if not rows:
        return None

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""

            # Parse inline formatting for cell content
            parts = parse_inline_formatting(cell_text)
            first = True
            for part_type, part_text in parts:
                if first:
                    run = cell.paragraphs[0].add_run(part_text)
                    first = False
                else:
                    run = cell.add_paragraph().add_run(part_text)
                run.font.name = "Calibri"
                run.font.size = Pt(10)

                if part_type == "bold":
                    run.bold = True
                elif part_type == "code":
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

            # Header row styling
            if row_idx == 0:
                set_cell_shading(cell, "D0D0D0")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    return table


def convert_md_to_docx(input_path, output_path):
    """Main conversion function."""
    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Read markdown file
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_block_lang = ""

    # Track counts
    paragraph_count = 0
    table_count = 0

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lang = line.strip()[3:].strip()
                code_block_lines = []
            else:
                # End code block - add as styled paragraph
                in_code_block = False
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                paragraph_count += 1

                # Add shading
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
                pPr.append(shd)

                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)

            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            style_name = f"Heading {level}"
            p = doc.add_heading(text, level=level)
            paragraph_count += 1

            # Ensure heading uses Calibri
            for run in p.runs:
                run.font.name = "Calibri"

            i += 1
            continue

        # Checkbox list items
        checkbox_match = re.match(r'^(\s*)- \[([ x])\]\s+(.+)$', line)
        if checkbox_match:
            indent = checkbox_match.group(1)
            checked = checkbox_match.group(2) == 'x'
            text = checkbox_match.group(3)

            p = doc.add_paragraph(style="List Bullet")
            paragraph_count += 1

            # Add checkbox symbol
            checkbox_char = "☐" if not checked else "☑"
            run = p.add_run(f"{checkbox_char} ")
            run.font.name = "Calibri"
            run.font.size = Pt(11)

            # Add text with formatting
            parts = parse_inline_formatting(text)
            for part_type, part_text in parts:
                run = p.add_run(part_text)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                if part_type == "bold":
                    run.bold = True
                elif part_type == "code":
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)

            i += 1
            continue

        # Bullet list items (simple -)
        bullet_match = re.match(r'^(\s*)-\s+(.+)$', line)
        if bullet_match:
            text = bullet_match.group(2)

            # But not if it's a checkbox (already handled above)
            if not text.strip().startswith('[ ]') and not text.strip().startswith('[x]'):
                p = doc.add_paragraph(style="List Bullet")
                paragraph_count += 1

                parts = parse_inline_formatting(text)
                for part_idx, (part_type, part_text) in enumerate(parts):
                    run = p.add_run(part_text)
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    if part_type == "bold":
                        run.bold = True
                    elif part_type == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)

                i += 1
                continue

        # Tables
        if line.strip().startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse and add table
            rows = parse_table_md('\n'.join(table_lines))
            if rows:
                add_table_to_doc(doc, rows)
                table_count += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="808080"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            paragraph_count += 1
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            add_formatted_paragraph(doc, line.strip())
            paragraph_count += 1

        i += 1

    # Save document
    doc.save(output_path)

    return paragraph_count, table_count


if __name__ == "__main__":
    print(f"Converting: {INPUT_FILE}")
    print(f"Output: {OUTPUT_FILE}")

    para_count, table_count = convert_md_to_docx(INPUT_FILE, OUTPUT_FILE)

    # Verify file exists and get size
    import os
    if os.path.exists(OUTPUT_FILE):
        file_size = os.path.getsize(OUTPUT_FILE)
        print(f"\n✓ Conversion successful!")
        print(f"  File: {OUTPUT_FILE}")
        print(f"  Size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
        print(f"  Paragraphs: {para_count}")
        print(f"  Tables: {table_count}")
    else:
        print("ERROR: File was not created!")